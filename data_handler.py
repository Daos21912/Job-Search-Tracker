import pandas as pd
import os
import openpyxl


DATA_FILE = "data/applications.xlsx"
COLUMNS = ["Company", "Role", "Date Applied", "Status", "Link", "Notes", "Last Updated"]

def init_data():
    os.makedirs("data", exist_ok=True)
    if not os.path.exists(DATA_FILE):
        df = pd.DataFrame(columns=COLUMNS)
        df.to_excel(DATA_FILE, index=False)

def load_data():
    init_data()
    df = pd.read_excel(DATA_FILE)
    if "Last Updated" not in df.columns:
        df["Last Updated"] = pd.Timestamp.today().normalize()
        save_data(df)
    df["Last Updated"] = pd.to_datetime(df["Last Updated"])
    df["Date Applied"] = pd.to_datetime(df["Date Applied"])
    # Fill any NaT values that slipped through
    df["Last Updated"] = df["Last Updated"].fillna(pd.Timestamp.today().normalize())
    save_data(df)
    return df

def save_data(df):
    df.to_excel(DATA_FILE, index=False)

def add_application(new_row):
    df = load_data()
    # Check for duplicate (exact match)
    duplicate = df[
        (df["Company"].str.strip().str.lower() == new_row["Company"].lower()) &
        (df["Role"].str.strip().str.lower() == new_row["Role"].lower())
    ]
    if not duplicate.empty:
        return False  # duplicate found

    new_row["Last Updated"] = pd.Timestamp.today().normalize()  # date only, no time
    df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    save_data(df)
    return True

def delete_application(index):
    df = load_data()
    df = df.drop(index).reset_index(drop=True)
    save_data(df)

CL_FILE = "data/cover_letters"
CL_SHEET = "Cover Letters"
CL_COLUMNS = ["Application Index", "Company", "Role", "Document Name", "Date Created", "File Path"]

def init_cover_letters():
    """Add Cover Letters sheet to xlsx if it doesn't exist, and create folder."""
    os.makedirs(CL_FILE, exist_ok=True)
    wb = openpyxl.load_workbook(DATA_FILE)
    if CL_SHEET not in wb.sheetnames:
        ws = wb.create_sheet(CL_SHEET)
        ws.append(CL_COLUMNS)
        wb.save(DATA_FILE)

def load_cover_letters():
    init_cover_letters()
    try:
        return pd.read_excel(DATA_FILE, sheet_name=CL_SHEET)
    except Exception:
        return pd.DataFrame(columns=CL_COLUMNS)

def save_cover_letter(app_index, company, role, doc_name, text):
    """Save cover letter as docx and log it in the xlsx."""
    from docx import Document
    
    os.makedirs(CL_FILE, exist_ok=True)
    
    # Build filename and path
    safe_name = doc_name.strip().replace(" ", "_")
    filename = f"{safe_name}.docx"
    filepath = os.path.join(CL_FILE, filename)
    
    # Create docx
    doc = Document()
    doc.add_heading(doc_name, level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filepath)
    
    # Log to xlsx
    init_cover_letters()
    wb = openpyxl.load_workbook(DATA_FILE)
    ws = wb[CL_SHEET]
    ws.append([
        app_index,
        company,
        role,
        doc_name.strip(),
        pd.Timestamp.today().normalize(),
        filepath
    ])
    wb.save(DATA_FILE)
    
    return filepath

def delete_cover_letter(record_index, filepath):
    """Remove the docx file and delete the record from the xlsx."""
    # Delete file if it exists
    if os.path.exists(filepath):
        os.remove(filepath)

    # Remove row from Cover Letters sheet
    wb = openpyxl.load_workbook(DATA_FILE)
    ws = wb[CL_SHEET]

    # Row 1 is headers, data starts at row 2, so xlsx row = record_index + 2
    ws.delete_rows(record_index + 2)
    wb.save(DATA_FILE)


def load_cover_letter_text(filepath):
    """Read text back out of a docx for editing."""
    from docx import Document
    if not os.path.exists(filepath):
        return ""
    doc = Document(filepath)
    # Skip the heading (first paragraph), join the rest
    paragraphs = [p.text for p in doc.paragraphs[1:]]
    return "\n".join(paragraphs)